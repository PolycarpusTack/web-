"""
Message Export API

This module provides comprehensive message and conversation export functionality
supporting multiple formats: JSON, CSV, PDF, DOCX, Markdown, and HTML.
"""
import asyncio
import json
import csv
import io
import base64
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import tempfile
import zipfile

from fastapi import APIRouter, Depends, HTTPException, Query, Response
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, or_, desc
from sqlalchemy.orm import selectinload
from pydantic import BaseModel, Field

from db.database import get_db
from db.models import User, Conversation, Message, Model
from auth.jwt import get_current_user

# Optional dependencies for advanced export formats
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import markdown
    from markdown.extensions import codehilite, tables, toc
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

# Pydantic Models
class ExportRequest(BaseModel):
    conversation_ids: List[str] = Field(..., description="List of conversation IDs to export")
    format: str = Field(..., description="Export format: json, csv, pdf, docx, md, html")
    include_metadata: bool = Field(True, description="Include message metadata")
    include_system_messages: bool = Field(False, description="Include system messages")
    include_timestamps: bool = Field(True, description="Include message timestamps")
    include_user_info: bool = Field(False, description="Include user information")
    include_model_info: bool = Field(True, description="Include model information")
    date_from: Optional[datetime] = Field(None, description="Filter messages from this date")
    date_to: Optional[datetime] = Field(None, description="Filter messages to this date")
    compress: bool = Field(False, description="Compress output as ZIP")
    custom_template: Optional[str] = Field(None, description="Custom template for formatting")

class BulkExportRequest(BaseModel):
    format: str = Field(..., description="Export format")
    folder_id: Optional[str] = Field(None, description="Export all conversations from folder")
    filter_shared: bool = Field(False, description="Only export shared conversations")
    filter_bookmarked: bool = Field(False, description="Only export bookmarked conversations")
    max_conversations: int = Field(100, description="Maximum conversations to export")

class ExportStatus(BaseModel):
    export_id: str
    status: str  # pending, processing, completed, failed
    progress: int  # 0-100
    created_at: datetime
    completed_at: Optional[datetime]
    download_url: Optional[str]
    file_size: Optional[int]
    error_message: Optional[str]

# Business Logic Classes
class MessageExporter:
    def __init__(self, db: AsyncSession):
        self.db = db
    
    async def export_conversations(self, user: User, request: ExportRequest) -> Union[bytes, str]:
        """Export conversations in the specified format."""
        # Get conversations with messages
        conversations = await self._get_conversations_with_messages(user, request)
        
        if not conversations:
            raise HTTPException(404, "No conversations found or access denied")
        
        # Export based on format
        if request.format.lower() == 'json':
            return await self._export_json(conversations, request)
        elif request.format.lower() == 'csv':
            return await self._export_csv(conversations, request)
        elif request.format.lower() == 'pdf':
            return await self._export_pdf(conversations, request)
        elif request.format.lower() == 'docx':
            return await self._export_docx(conversations, request)
        elif request.format.lower() == 'md':
            return await self._export_markdown(conversations, request)
        elif request.format.lower() == 'html':
            return await self._export_html(conversations, request)
        else:
            raise HTTPException(400, f"Unsupported export format: {request.format}")
    
    async def _get_conversations_with_messages(self, user: User, request: ExportRequest) -> List[Conversation]:
        """Get conversations with their messages."""
        query = select(Conversation).options(
            selectinload(Conversation.messages).selectinload(Message.user),
            selectinload(Conversation.model),
            selectinload(Conversation.users)
        ).where(
            and_(
                Conversation.id.in_(request.conversation_ids),
                Conversation.users.any(User.id == user.id)
            )
        )
        
        result = await self.db.execute(query)
        conversations = result.scalars().all()
        
        # Filter messages by date if specified
        if request.date_from or request.date_to:
            for conversation in conversations:
                filtered_messages = []
                for message in conversation.messages:
                    if request.date_from and message.created_at < request.date_from:
                        continue
                    if request.date_to and message.created_at > request.date_to:
                        continue
                    filtered_messages.append(message)
                conversation.messages = filtered_messages
        
        # Filter system messages if not included
        if not request.include_system_messages:
            for conversation in conversations:
                conversation.messages = [m for m in conversation.messages if m.role != 'system']
        
        return conversations
    
    async def _export_json(self, conversations: List[Conversation], request: ExportRequest) -> str:
        """Export as JSON format."""
        export_data = {
            "export_info": {
                "format": "json",
                "exported_at": datetime.utcnow().isoformat(),
                "conversation_count": len(conversations),
                "include_metadata": request.include_metadata,
                "include_timestamps": request.include_timestamps
            },
            "conversations": []
        }
        
        for conversation in conversations:
            conv_data = {
                "id": conversation.id,
                "title": conversation.title,
                "created_at": conversation.created_at.isoformat() if request.include_timestamps else None,
                "updated_at": conversation.updated_at.isoformat() if request.include_timestamps else None,
                "model": {
                    "id": conversation.model.id,
                    "name": conversation.model.name,
                    "provider": conversation.model.provider
                } if request.include_model_info and conversation.model else None,
                "system_prompt": conversation.system_prompt if request.include_metadata else None,
                "messages": []
            }
            
            for message in conversation.messages:
                msg_data = {
                    "id": message.id,
                    "role": message.role,
                    "content": message.content,
                    "created_at": message.created_at.isoformat() if request.include_timestamps else None
                }
                
                if request.include_metadata:
                    msg_data.update({
                        "tokens": message.tokens,
                        "cost": message.cost,
                        "metadata": message.meta_data
                    })
                
                if request.include_user_info and message.user:
                    msg_data["user"] = {
                        "id": message.user.id,
                        "username": message.user.username
                    }
                
                conv_data["messages"].append(msg_data)
            
            export_data["conversations"].append(conv_data)
        
        return json.dumps(export_data, indent=2, ensure_ascii=False)
    
    async def _export_csv(self, conversations: List[Conversation], request: ExportRequest) -> str:
        """Export as CSV format."""
        output = io.StringIO()
        
        # Define CSV headers
        headers = ["conversation_id", "conversation_title", "message_id", "role", "content"]
        
        if request.include_timestamps:
            headers.extend(["message_created_at", "conversation_created_at"])
        
        if request.include_model_info:
            headers.extend(["model_id", "model_name", "model_provider"])
        
        if request.include_metadata:
            headers.extend(["tokens", "cost"])
        
        if request.include_user_info:
            headers.extend(["user_id", "username"])
        
        writer = csv.writer(output)
        writer.writerow(headers)
        
        # Write data
        for conversation in conversations:
            for message in conversation.messages:
                row = [
                    conversation.id,
                    conversation.title,
                    message.id,
                    message.role,
                    message.content
                ]
                
                if request.include_timestamps:
                    row.extend([
                        message.created_at.isoformat(),
                        conversation.created_at.isoformat()
                    ])
                
                if request.include_model_info and conversation.model:
                    row.extend([
                        conversation.model.id,
                        conversation.model.name,
                        conversation.model.provider
                    ])
                elif request.include_model_info:
                    row.extend(["", "", ""])
                
                if request.include_metadata:
                    row.extend([
                        message.tokens or 0,
                        message.cost or 0.0
                    ])
                
                if request.include_user_info:
                    if message.user:
                        row.extend([message.user.id, message.user.username])
                    else:
                        row.extend(["", ""])
                
                writer.writerow(row)
        
        return output.getvalue()
    
    async def _export_pdf(self, conversations: List[Conversation], request: ExportRequest) -> bytes:
        """Export as PDF format."""
        if not HAS_REPORTLAB:
            raise HTTPException(400, "PDF export requires reportlab package")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph("Conversation Export", title_style))
        story.append(Spacer(1, 20))
        
        # Export info
        export_info = f"Exported on: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}<br/>"
        export_info += f"Conversations: {len(conversations)}<br/>"
        export_info += f"Format: PDF"
        story.append(Paragraph(export_info, styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Conversations
        for i, conversation in enumerate(conversations):
            if i > 0:
                story.append(Spacer(1, 30))
            
            # Conversation header
            conv_title = f"Conversation: {conversation.title}"
            story.append(Paragraph(conv_title, styles['Heading1']))
            
            if request.include_model_info and conversation.model:
                model_info = f"Model: {conversation.model.name} ({conversation.model.provider})"
                story.append(Paragraph(model_info, styles['Normal']))
            
            if request.include_timestamps:
                date_info = f"Created: {conversation.created_at.strftime('%Y-%m-%d %H:%M:%S')}"
                story.append(Paragraph(date_info, styles['Normal']))
            
            story.append(Spacer(1, 15))
            
            # Messages
            for message in conversation.messages:
                # Message header
                role_style = styles['Heading3'] if message.role == 'assistant' else styles['Heading4']
                role_text = message.role.title()
                
                if request.include_user_info and message.user:
                    role_text += f" ({message.user.username})"
                
                if request.include_timestamps:
                    role_text += f" - {message.created_at.strftime('%H:%M:%S')}"
                
                story.append(Paragraph(role_text, role_style))
                
                # Message content
                content = message.content.replace('\n', '<br/>')
                story.append(Paragraph(content, styles['Normal']))
                
                if request.include_metadata and (message.tokens or message.cost):
                    meta_text = f"Tokens: {message.tokens or 0}, Cost: ${message.cost or 0:.4f}"
                    story.append(Paragraph(meta_text, styles['Italic']))
                
                story.append(Spacer(1, 10))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def _export_docx(self, conversations: List[Conversation], request: ExportRequest) -> bytes:
        """Export as DOCX format."""
        if not HAS_DOCX:
            raise HTTPException(400, "DOCX export requires python-docx package")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Conversation Export', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Export info
        info_p = doc.add_paragraph()
        info_p.add_run(f"Exported on: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
        info_p.add_run(f"Conversations: {len(conversations)}\n")
        info_p.add_run(f"Format: Microsoft Word Document")
        
        # Conversations
        for conversation in conversations:
            doc.add_page_break()
            
            # Conversation title
            doc.add_heading(f"Conversation: {conversation.title}", 1)
            
            # Conversation metadata
            if request.include_model_info and conversation.model:
                model_p = doc.add_paragraph()
                model_p.add_run("Model: ").bold = True
                model_p.add_run(f"{conversation.model.name} ({conversation.model.provider})")
            
            if request.include_timestamps:
                date_p = doc.add_paragraph()
                date_p.add_run("Created: ").bold = True
                date_p.add_run(conversation.created_at.strftime('%Y-%m-%d %H:%M:%S'))
            
            if conversation.system_prompt and request.include_metadata:
                sys_p = doc.add_paragraph()
                sys_p.add_run("System Prompt: ").bold = True
                sys_p.add_run(conversation.system_prompt)
            
            # Messages
            for message in conversation.messages:
                # Message header
                msg_header = doc.add_heading(level=3)
                role_text = message.role.title()
                
                if request.include_user_info and message.user:
                    role_text += f" ({message.user.username})"
                
                if request.include_timestamps:
                    role_text += f" - {message.created_at.strftime('%H:%M:%S')}"
                
                msg_header.add_run(role_text)
                
                # Message content
                content_p = doc.add_paragraph(message.content)
                
                # Message metadata
                if request.include_metadata and (message.tokens or message.cost):
                    meta_p = doc.add_paragraph()
                    meta_run = meta_p.add_run(f"Tokens: {message.tokens or 0}, Cost: ${message.cost or 0:.4f}")
                    meta_run.italic = True
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def _export_markdown(self, conversations: List[Conversation], request: ExportRequest) -> str:
        """Export as Markdown format."""
        lines = []
        
        # Title
        lines.append("# Conversation Export")
        lines.append("")
        
        # Export info
        lines.append(f"**Exported on:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"**Conversations:** {len(conversations)}")
        lines.append(f"**Format:** Markdown")
        lines.append("")
        
        # Table of contents
        lines.append("## Table of Contents")
        lines.append("")
        for i, conversation in enumerate(conversations, 1):
            title = conversation.title.replace('#', '\\#')  # Escape hash symbols
            lines.append(f"{i}. [{title}](#conversation-{i})")
        lines.append("")
        
        # Conversations
        for i, conversation in enumerate(conversations, 1):
            lines.append("---")
            lines.append("")
            lines.append(f"## Conversation {i}: {conversation.title} {{#conversation-{i}}}")
            lines.append("")
            
            # Conversation metadata
            if request.include_model_info and conversation.model:
                lines.append(f"**Model:** {conversation.model.name} ({conversation.model.provider})")
            
            if request.include_timestamps:
                lines.append(f"**Created:** {conversation.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            
            if conversation.system_prompt and request.include_metadata:
                lines.append(f"**System Prompt:** {conversation.system_prompt}")
            
            lines.append("")
            
            # Messages
            for message in conversation.messages:
                # Message header
                role_emoji = {
                    'user': '👤',
                    'assistant': '🤖',
                    'system': '⚙️'
                }.get(message.role, '💬')
                
                header = f"### {role_emoji} {message.role.title()}"
                
                if request.include_user_info and message.user:
                    header += f" ({message.user.username})"
                
                if request.include_timestamps:
                    header += f" - {message.created_at.strftime('%H:%M:%S')}"
                
                lines.append(header)
                lines.append("")
                
                # Message content
                lines.append(message.content)
                lines.append("")
                
                # Message metadata
                if request.include_metadata and (message.tokens or message.cost):
                    lines.append(f"*Tokens: {message.tokens or 0}, Cost: ${message.cost or 0:.4f}*")
                    lines.append("")
        
        return "\n".join(lines)
    
    async def _export_html(self, conversations: List[Conversation], request: ExportRequest) -> str:
        """Export as HTML format."""
        html_parts = []
        
        # HTML header
        html_parts.append("""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Conversation Export</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .header { text-align: center; margin-bottom: 40px; }
        .conversation { margin-bottom: 60px; border-bottom: 2px solid #eee; padding-bottom: 40px; }
        .conversation:last-child { border-bottom: none; }
        .conversation-title { color: #333; border-bottom: 1px solid #ddd; padding-bottom: 10px; }
        .conversation-meta { color: #666; font-size: 14px; margin-bottom: 20px; }
        .message { margin-bottom: 20px; padding: 15px; border-radius: 8px; }
        .message.user { background-color: #e3f2fd; border-left: 4px solid #2196f3; }
        .message.assistant { background-color: #f3e5f5; border-left: 4px solid #9c27b0; }
        .message.system { background-color: #fff3e0; border-left: 4px solid #ff9800; }
        .message-header { font-weight: bold; margin-bottom: 8px; color: #555; }
        .message-content { white-space: pre-wrap; }
        .message-meta { font-size: 12px; color: #888; margin-top: 8px; font-style: italic; }
        .toc { margin-bottom: 40px; }
        .toc ul { list-style-type: none; padding-left: 0; }
        .toc li { margin-bottom: 5px; }
        .toc a { text-decoration: none; color: #2196f3; }
        .toc a:hover { text-decoration: underline; }
    </style>
</head>
<body>
        """)
        
        # Header
        html_parts.append(f"""
    <div class="header">
        <h1>Conversation Export</h1>
        <p>Exported on: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p>Conversations: {len(conversations)} | Format: HTML</p>
    </div>
        """)
        
        # Table of contents
        html_parts.append('<div class="toc"><h2>Table of Contents</h2><ul>')
        for i, conversation in enumerate(conversations, 1):
            title = conversation.title.replace('<', '&lt;').replace('>', '&gt;')
            html_parts.append(f'<li><a href="#conversation-{i}">{i}. {title}</a></li>')
        html_parts.append('</ul></div>')
        
        # Conversations
        for i, conversation in enumerate(conversations, 1):
            title = conversation.title.replace('<', '&lt;').replace('>', '&gt;')
            html_parts.append(f'<div class="conversation" id="conversation-{i}">')
            html_parts.append(f'<h2 class="conversation-title">Conversation {i}: {title}</h2>')
            
            # Conversation metadata
            meta_parts = []
            if request.include_model_info and conversation.model:
                meta_parts.append(f"Model: {conversation.model.name} ({conversation.model.provider})")
            if request.include_timestamps:
                meta_parts.append(f"Created: {conversation.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            
            if meta_parts:
                html_parts.append(f'<div class="conversation-meta">{" | ".join(meta_parts)}</div>')
            
            # System prompt
            if conversation.system_prompt and request.include_metadata:
                sys_prompt = conversation.system_prompt.replace('<', '&lt;').replace('>', '&gt;')
                html_parts.append(f'<div class="message system"><div class="message-header">System Prompt</div><div class="message-content">{sys_prompt}</div></div>')
            
            # Messages
            for message in conversation.messages:
                content = message.content.replace('<', '&lt;').replace('>', '&gt;')
                
                # Message header
                header_parts = [message.role.title()]
                if request.include_user_info and message.user:
                    header_parts.append(f"({message.user.username})")
                if request.include_timestamps:
                    header_parts.append(f"- {message.created_at.strftime('%H:%M:%S')}")
                
                header = " ".join(header_parts)
                
                html_parts.append(f'<div class="message {message.role}">')
                html_parts.append(f'<div class="message-header">{header}</div>')
                html_parts.append(f'<div class="message-content">{content}</div>')
                
                # Message metadata
                if request.include_metadata and (message.tokens or message.cost):
                    meta_text = f"Tokens: {message.tokens or 0}, Cost: ${message.cost or 0:.4f}"
                    html_parts.append(f'<div class="message-meta">{meta_text}</div>')
                
                html_parts.append('</div>')
            
            html_parts.append('</div>')
        
        # HTML footer
        html_parts.append('</body></html>')
        
        return "\n".join(html_parts)


# Router
router = APIRouter(prefix="/api/export", tags=["Message Export"])

@router.post("/conversations")
async def export_conversations(
    request: ExportRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Export conversations in the specified format."""
    exporter = MessageExporter(db)
    
    try:
        export_data = await exporter.export_conversations(current_user, request)
        
        # Determine content type and filename
        content_types = {
            'json': 'application/json',
            'csv': 'text/csv',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'md': 'text/markdown',
            'html': 'text/html'
        }
        
        file_extensions = {
            'json': 'json',
            'csv': 'csv', 
            'pdf': 'pdf',
            'docx': 'docx',
            'md': 'md',
            'html': 'html'
        }
        
        content_type = content_types.get(request.format.lower(), 'application/octet-stream')
        extension = file_extensions.get(request.format.lower(), 'txt')
        
        # Generate filename
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        conv_count = len(request.conversation_ids)
        filename = f"conversations_export_{conv_count}_{timestamp}.{extension}"
        
        # Create response
        if isinstance(export_data, bytes):
            return Response(
                content=export_data,
                media_type=content_type,
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        else:
            return Response(
                content=export_data.encode('utf-8'),
                media_type=content_type,
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
    
    except Exception as e:
        raise HTTPException(500, f"Export failed: {str(e)}")

@router.get("/formats")
async def get_supported_formats():
    """Get list of supported export formats with their capabilities."""
    formats = {
        "json": {
            "name": "JSON",
            "description": "Structured data format, preserves all metadata",
            "supports_metadata": True,
            "supports_formatting": False,
            "file_extension": "json",
            "available": True
        },
        "csv": {
            "name": "CSV",
            "description": "Spreadsheet format, good for analysis",
            "supports_metadata": True,
            "supports_formatting": False,
            "file_extension": "csv",
            "available": True
        },
        "pdf": {
            "name": "PDF",
            "description": "Professional document format",
            "supports_metadata": True,
            "supports_formatting": True,
            "file_extension": "pdf",
            "available": HAS_REPORTLAB
        },
        "docx": {
            "name": "Word Document",
            "description": "Microsoft Word format",
            "supports_metadata": True,
            "supports_formatting": True,
            "file_extension": "docx",
            "available": HAS_DOCX
        },
        "md": {
            "name": "Markdown",
            "description": "Text format with markup",
            "supports_metadata": True,
            "supports_formatting": True,
            "file_extension": "md",
            "available": True
        },
        "html": {
            "name": "HTML",
            "description": "Web page format",
            "supports_metadata": True,
            "supports_formatting": True,
            "file_extension": "html",
            "available": True
        }
    }
    
    return {
        "supported_formats": formats,
        "recommended_packages": {
            "pdf": "reportlab" if not HAS_REPORTLAB else None,
            "docx": "python-docx" if not HAS_DOCX else None,
            "advanced_markdown": "markdown" if not HAS_MARKDOWN else None
        }
    }